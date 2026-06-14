from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color):
    properties = cell._element.tcPr
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    properties.append(shading)

def create_diet_plan_doc(patient, diet_plan, output_path):
    doc = Document()
    
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    clinic_name = header.add_run("NUTRISLIMS HEALTH & WELLNESS CLINIC\n")
    clinic_name.bold = True
    clinic_name.font.size = Pt(16)
    clinic_name.font.color.rgb = RGBColor(74, 144, 226)
    clinic_name.font.name = 'Arial'
    
    location = header.add_run("Indore, Madhya Pradesh\n\n")
    location.font.size = Pt(11)
    location.font.name = 'Arial'
    
    dr_info = header.add_run(
        "Dr. Oshin Ambekar\n"
        "B.Sc. Nutrition | M.Sc. Clinical Nutrition\n"
        "Doctor of Philosophy (PhD) in Clinical Nutrition\n"
        "Certified Clinical Dietitian | Certified Weight Management Specialist\n"
    )
    dr_info.font.size = Pt(10)
    dr_info.font.name = 'Arial'
    
    doc.add_paragraph("_" * 65).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title = doc.add_paragraph("Personalized Clinical Diet Plan Report")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(74, 144, 226)
    title_run.font.name = 'Arial'
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph("_" * 65).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(74, 144, 226)
        run.font.name = 'Arial'
        return p
        
    def format_table(table):
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
    add_heading("Patient Profile")
    pt_table = doc.add_table(rows=4, cols=2)
    
    details = [
        ("Name", patient.get('name', 'N/A')),
        ("Age & Gender", f"{patient.get('age', 'N/A')} / {patient.get('gender', 'N/A')}"),
        ("Weight & Height", f"{patient.get('weight', 'N/A')} / {patient.get('height', 'N/A')}"),
        ("BMI & Conditions", f"{patient.get('bmi', 'N/A')} ({patient.get('bmi_category', 'N/A')}) - {', '.join(patient.get('conditions', [])) if patient.get('conditions') else 'None'}")
    ]
    
    for i, (key, value) in enumerate(details):
        pt_table.rows[i].cells[0].text = key
        pt_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_background(pt_table.rows[i].cells[0], "E8F1FA")
        pt_table.rows[i].cells[1].text = str(value)
        
    format_table(pt_table)
    
    add_heading("Prescribed Diet Plan")
    
    for meal in diet_plan.get('meals', []):
        p = doc.add_paragraph()
        run = p.add_run(meal.get('mealName', 'Meal'))
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        
        meal_table = doc.add_table(rows=8, cols=3)
        meal_table.rows[0].cells[0].text = "Option"
        meal_table.rows[0].cells[1].text = "Food Item / Recipe"
        meal_table.rows[0].cells[2].text = "Quantity / Notes"
        
        for cell in meal_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_background(cell, "E8F1FA")
            
        options = meal.get('options', [])
        for i in range(7):
            opt_text = options[i] if i < len(options) else ""
            meal_table.rows[i+1].cells[0].text = f"Option {i+1}"
            meal_table.rows[i+1].cells[0].paragraphs[0].runs[0].bold = True
            meal_table.rows[i+1].cells[1].text = opt_text
            meal_table.rows[i+1].cells[2].text = "As prescribed"
            
        format_table(meal_table)
        
    add_heading("Dietitian Guidelines & Notes")
    notes = diet_plan.get('dietitianNotes', [])
    notes_table = doc.add_table(rows=len(notes)+1, cols=2)
    notes_table.rows[0].cells[0].text = "#"
    notes_table.rows[0].cells[1].text = "Guideline"
    set_cell_background(notes_table.rows[0].cells[0], "E8F1FA")
    set_cell_background(notes_table.rows[0].cells[1], "E8F1FA")
    
    for i, note in enumerate(notes):
        notes_table.rows[i+1].cells[0].text = str(i+1)
        notes_table.rows[i+1].cells[1].text = note
    format_table(notes_table)
    
    add_heading("Foods to Strictly Avoid")
    avoids = diet_plan.get('foodsToAvoid', [])
    while len(avoids) < 5:
        avoids.append("None")
        
    avoid_table = doc.add_table(rows=len(avoids)+1, cols=1)
    avoid_table.rows[0].cells[0].text = "Restricted Items"
    avoid_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(avoid_table.rows[0].cells[0], "FADCDA")
    
    for i, avoid in enumerate(avoids):
        avoid_table.rows[i+1].cells[0].text = f"❌ {avoid}"
        
    format_table(avoid_table)
    
    doc.save(output_path)
